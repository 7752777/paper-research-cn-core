from __future__ import annotations

import argparse
import json
import re
from pathlib import Path


def _table_rows(markdown: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in markdown:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and not all(set(cell) <= {"-", ":"} for cell in cells):
            rows.append(cells)
    return rows


def _clean_text(text: str) -> str:
    text = text.replace("**", "")
    labels = {
        "MiakeLye2016": "Miake-Lye et al., 2016",
        "ArkseyOmalley2005": "Arksey and O'Malley, 2005",
        "Tricco2018": "Tricco et al., 2018",
    }

    def citation(match: re.Match[str]) -> str:
        keys = [key.strip().lstrip("@") for key in match.group(1).split(";")]
        return "（" + "；".join(labels.get(key, key) for key in keys) + "）"

    return re.sub(r"\[([^\]]*@[^\]]+)\]", citation, text)


def _add_table(doc, table_data: dict[str, object]) -> None:
    if table_data.get("caption"):
        doc.add_paragraph(_clean_text(str(table_data["caption"])))
    rows = _table_rows(list(table_data.get("markdown", [])))
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, cells in enumerate(rows):
        for column_index, cell in enumerate(cells):
            if column_index < len(table.rows[row_index].cells):
                table.rows[row_index].cells[column_index].text = _clean_text(cell)


def render(document: dict[str, object], output: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as error:
        raise RuntimeError("python-docx is required to render DOCX") from error
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(str(document["title"]))
    run.bold = True
    run.font.name = "黑体"
    run.font.size = Pt(16)
    source_parent = Path(str(document.get("source", output))).parent
    blocks = document.get("blocks")
    if not blocks:
        blocks = [{"type": "heading", "level": section.get("level", 1), "text": section.get("title", "")} for section in document.get("sections", [])]
    for block in blocks:
        kind = block.get("type")
        if kind == "heading":
            doc.add_heading(_clean_text(str(block.get("text", ""))), level=min(int(block.get("level", 1)), 3))
        elif kind == "paragraph":
            doc.add_paragraph(_clean_text(str(block.get("text", ""))))
        elif kind == "table":
            _add_table(doc, block)
        elif kind == "figure":
            image = block.get("path")
            if image:
                image_path = source_parent / str(image)
                if image_path.is_file():
                    doc.add_picture(str(image_path), width=Inches(6.0))
            caption = str(block.get("caption", ""))
            if caption:
                paragraph = doc.add_paragraph(_clean_text(caption))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Render a doc.json contract as a Chinese-journal DOCX draft.")
    parser.add_argument("doc_json", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    try:
        render(json.loads(args.doc_json.read_text(encoding="utf-8")), args.output)
    except RuntimeError as error:
        print(str(error))
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
