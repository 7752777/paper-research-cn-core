from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "skills" / "paper-research-cn-core" / "scripts"))


class PackageContractTests(unittest.TestCase):
    def test_exactly_one_skill_is_discoverable_and_valid(self) -> None:
        from install import discover_skills
        from validate_package import validate_package

        self.assertEqual(discover_skills(ROOT), ["paper-research-cn-core"])
        self.assertEqual(validate_package(ROOT), [])

    def test_ai_style_audit_blocks_submission_meta_language(self) -> None:
        from audit_ai_style import audit_text

        findings = audit_text("本文已按目标期刊投稿要求完成，等待评审。")
        self.assertEqual(findings[0]["rule_id"], "AI-STYLE-001")
        self.assertEqual(findings[0]["severity"], "major")
        self.assertEqual(audit_text("本文讨论某导航系统在应急通信中的作用边界。"), [])

    def test_reference_audit_requires_complete_journal_fields(self) -> None:
        from audit_refs import audit_bibliography

        findings = audit_bibliography("[1] 张三. 导航应急研究[J]. 某期刊, 2024.\n")
        self.assertEqual(findings[0]["rule_id"], "REF-GBT-001")
        self.assertEqual(findings[0]["severity"], "major")

    def test_reference_audit_requires_official_metadata_status(self) -> None:
        from audit_refs import audit_bibliography

        entry = "[1] Demo A. Verified study[J]. Demo Journal, 2024, 1(1): 1-2. DOI: 10.1000/demo.\n"
        missing = audit_bibliography(entry)
        self.assertIn("REF-METADATA-001", {item["rule_id"] for item in missing})
        verified = audit_bibliography(entry, metadata_records=[{
            "doi": "10.1000/demo",
            "metadata_status": "official_verified",
            "verification_source": "publisher metadata",
        }])
        self.assertEqual(verified, [])

    def test_figure_audit_requires_denominator_and_vector_source(self) -> None:
        from audit_figures import audit_markdown

        findings = audit_markdown(
            "图1 主题分布。\n\n![主题分布](figures/theme.png)\n",
            ROOT,
        )
        rule_ids = {item["rule_id"] for item in findings}
        self.assertIn("FIG-CAPTION-001", rule_ids)
        self.assertIn("FIG-VECTOR-001", rule_ids)

    def test_figure_audit_requires_caption_for_every_asset(self) -> None:
        from audit_figures import audit_markdown

        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            figures = root / "figures"
            figures.mkdir()
            (figures / "flow.svg").write_text("<svg/>", encoding="utf-8")
            findings = audit_markdown("![flow](figures/flow.svg)\n", root)
        self.assertIn("FIG-CAPTION-002", {item["rule_id"] for item in findings})

    def test_install_hashes_can_compare_a_complete_source_tree(self) -> None:
        from verify_install import compare_skill_trees, tree_hash

        source = ROOT / "skills" / "paper-research-cn-core"
        self.assertTrue(tree_hash(source))
        self.assertEqual(compare_skill_trees(source, source), [])

    def test_doc_json_retains_figure_asset_path(self) -> None:
        from build_doc_json import build_document

        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / "draft.qmd"
            source.write_text("---\ntitle: 测试\n---\n\n图1 测试（n=1）\n\n![测试](figures/test.png)\n", encoding="utf-8")
            document = build_document(source)
        self.assertEqual(document["figures"][0]["path"], "figures/test.png")
        self.assertNotIn("---", [block.get("text") for block in document["blocks"]])

    def test_doc_json_keeps_tables_and_figures_in_document_order(self) -> None:
        from build_doc_json import build_document

        with tempfile.TemporaryDirectory() as temporary:
            source = Path(temporary) / "draft.qmd"
            source.write_text("# 一、方法\n\n表1 样本（n=1）\n\n| A |\n| --- |\n| 1 |\n\n图1 流程（n=1）\n\n![流程](figures/flow.png)\n", encoding="utf-8")
            document = build_document(source)
        self.assertEqual([block["type"] for block in document["blocks"]], ["heading", "table", "figure"])

    def test_doc_json_materializes_declared_bibliography_at_reference_heading(self) -> None:
        from build_doc_json import build_document

        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            (root / "refs.bib").write_text("@article{Demo2024,\n author = {Demo A},\n title = {Demo title},\n journal = {Demo Journal},\n year = {2024},\n volume = {1},\n pages = {1--2}\n}\n", encoding="utf-8")
            source = root / "draft.qmd"
            source.write_text("---\ntitle: 测试\nbibliography: refs.bib\n---\n\n# 参考文献\n", encoding="utf-8")
            document = build_document(source)
        self.assertTrue(document["blocks"][-1]["text"].startswith("[1] Demo A."))

    def test_pdf_renderer_exposes_a_local_docx_fallback(self) -> None:
        from render_pdf import available_backends

        self.assertIsInstance(available_backends(), tuple)

    def test_pdf_renderer_reads_docx_paragraphs_and_tables_in_xml_order(self) -> None:
        from docx import Document
        from render_pdf import docx_block_types

        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "ordered.docx"
            document = Document()
            document.add_paragraph("前段")
            document.add_table(rows=1, cols=1).cell(0, 0).text = "表格"
            document.add_paragraph("后段")
            document.save(path)
            self.assertEqual(docx_block_types(path), ["paragraph", "table", "paragraph"])

    def test_pdf_renderer_keeps_a_figure_and_its_caption_together(self) -> None:
        from render_pdf import group_blocks

        groups = group_blocks([
            {"type": "image", "blob": b"", "suffix": ".png"},
            {"type": "paragraph", "text": "图1 流程图"},
        ])
        self.assertEqual([len(group) for group in groups], [2])


if __name__ == "__main__":
    unittest.main()
