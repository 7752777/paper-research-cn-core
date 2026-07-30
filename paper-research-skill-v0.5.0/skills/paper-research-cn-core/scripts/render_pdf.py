from __future__ import annotations

import argparse
import importlib.util
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path


FONT_ENV = "CN_CORE_CJK_FONT"


def fallback_font_path() -> Path | None:
    candidates: list[Path] = []
    if os.environ.get(FONT_ENV):
        candidates.append(Path(os.environ[FONT_ENV]))
    if sys.platform.startswith("win"):
        candidates.extend((
            Path(os.environ.get("WINDIR", r"C:\\Windows")) / "Fonts" / "simhei.ttf",
            Path(os.environ.get("WINDIR", r"C:\\Windows")) / "Fonts" / "msyh.ttf",
        ))
    elif sys.platform == "darwin":
        candidates.extend((
            Path("/Library/Fonts/NotoSansCJKsc-Regular.otf"),
            Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
        ))
    else:
        candidates.extend((
            Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
            Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttf"),
        ))
    return next((path for path in candidates if path.is_file()), None)


def fallback_diagnostics() -> list[str]:
    missing = [name for name in ("docx", "reportlab") if importlib.util.find_spec(name) is None]
    if fallback_font_path() is None:
        missing.append(f"a CJK font ({FONT_ENV} or a platform font)")
    return missing


def available_backends() -> tuple[str, ...]:
    backends: list[str] = []
    if shutil.which("soffice") or shutil.which("libreoffice"):
        backends.append("libreoffice")
    if not fallback_diagnostics():
        backends.append("reportlab-docx")
    return tuple(backends)


def _paragraph_style(font_path: Path):
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = "CNCoreCJK"
    if font_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
    styles = getSampleStyleSheet()
    body = ParagraphStyle("ChineseBody", parent=styles["BodyText"], fontName=font_name, fontSize=9.5, leading=15)
    title = ParagraphStyle("ChineseTitle", parent=styles["Title"], fontName=font_name, fontSize=16, leading=22, alignment=TA_CENTER)
    heading = ParagraphStyle("ChineseHeading", parent=styles["Heading2"], fontName=font_name, fontSize=12, leading=18)
    return body, title, heading


def _docx_blocks(docx_path: Path) -> list[dict[str, object]]:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    source = Document(docx_path)
    paragraphs = {id(paragraph._p): paragraph for paragraph in source.paragraphs}
    tables = {id(table._tbl): table for table in source.tables}
    blip_tag = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    embed_tag = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    blocks: list[dict[str, object]] = []
    for child in source.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = paragraphs.get(id(child)) or Paragraph(child, source)
            text = paragraph.text.strip()
            if text:
                kind = "heading" if paragraph.style.name.lower().startswith("heading") else "paragraph"
                blocks.append({"type": kind, "text": text})
            for blip in child.iter(blip_tag):
                relationship_id = blip.get(embed_tag)
                if relationship_id and relationship_id in source.part.related_parts:
                    image = source.part.related_parts[relationship_id]
                    blocks.append({"type": "image", "blob": image.blob, "suffix": Path(image.partname).suffix or ".png"})
        elif child.tag.endswith("}tbl"):
            table = tables.get(id(child)) or Table(child, source)
            blocks.append({"type": "table", "table": table})
    return blocks


def docx_block_types(docx_path: Path) -> list[str]:
    return [str(block["type"]) for block in _docx_blocks(docx_path)]


def group_blocks(blocks: list[dict[str, object]]) -> list[list[dict[str, object]]]:
    groups: list[list[dict[str, object]]] = []
    index = 0
    while index < len(blocks):
        current = blocks[index]
        if current["type"] == "image" and index + 1 < len(blocks):
            following = blocks[index + 1]
            if following["type"] == "paragraph" and str(following.get("text", "")).startswith("图"):
                groups.append([current, following])
                index += 2
                continue
        groups.append([current])
        index += 1
    return groups


def _fallback_render(docx_path: Path, pdf_path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import Image, KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_path = fallback_font_path()
    if font_path is None:
        raise RuntimeError(f"No CJK font is available for the local fallback; set {FONT_ENV} to a readable font path.")
    body, title, heading = _paragraph_style(font_path)
    story = []
    first = True
    with tempfile.TemporaryDirectory() as temporary:
        tmp = Path(temporary)
        for group_index, group in enumerate(group_blocks(_docx_blocks(docx_path))):
            flowables = []
            for block_index, block in enumerate(group):
                kind = block["type"]
                if kind in {"paragraph", "heading"}:
                    text = str(block["text"])
                    style = title if first else (heading if kind == "heading" else body)
                    first = False
                    flowables.extend((Paragraph(text, style), Spacer(1, 0.18 * cm)))
                elif kind == "table":
                    table = block["table"]
                    values = [[Paragraph(cell.text or "", body) for cell in row.cells] for row in table.rows]
                    if not values:
                        continue
                    rendered = Table(values, repeatRows=1, hAlign="LEFT")
                    rendered.setStyle(TableStyle([
                        ("FONTNAME", (0, 0), (-1, -1), "CNCoreCJK"),
                        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                        ("LEADING", (0, 0), (-1, -1), 12),
                        ("LINEABOVE", (0, 0), (-1, 0), 0.8, colors.black),
                        ("LINEBELOW", (0, 0), (-1, 0), 0.6, colors.black),
                        ("LINEBELOW", (0, -1), (-1, -1), 0.8, colors.black),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ]))
                    flowables.extend((rendered, Spacer(1, 0.28 * cm)))
                elif kind == "image":
                    image_path = tmp / f"image-{group_index}-{block_index}{block['suffix']}"
                    image_path.write_bytes(block["blob"])
                    flowables.extend((Image(str(image_path), width=15 * cm, height=8.2 * cm, kind="proportional"), Spacer(1, 0.22 * cm)))
            if len(group) > 1 and group[0]["type"] == "image":
                story.append(KeepTogether(flowables))
            else:
                story.extend(flowables)
        document = SimpleDocTemplate(str(pdf_path), pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm, topMargin=1.8 * cm, bottomMargin=1.8 * cm)
        document.build(story)


def convert(docx: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = output_dir / f"{docx.stem}.pdf"
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        result = subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx)], check=False)
        if result.returncode == 0 and pdf_path.is_file():
            return pdf_path
    missing = fallback_diagnostics()
    if missing:
        raise RuntimeError(f"Local DOCX/PDF fallback is unavailable: missing {', '.join(missing)}. Install requirements.txt and set {FONT_ENV} when needed.")
    _fallback_render(docx, pdf_path)
    return pdf_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert DOCX to PDF through LibreOffice when available, otherwise a local python-docx/reportlab fallback.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    try:
        print(convert(args.docx, args.output_dir))
    except Exception as error:
        print(f"DOCX-to-PDF conversion failed: {error}")
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
